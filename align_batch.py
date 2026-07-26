"""
align_batch.py — Batch forced-alignment engine for the IELTS audio-script marker workflow.

For every (audio, text) pair found in the input folders, this script:
  1. Runs word-level CTC forced alignment (ctc-forced-aligner / wav2vec2-MMS) between
     the audio and the ALREADY-CORRECT transcript text.
  2. Writes a marker-embedded .txt file using the exact "[ms]word[ms]word..." syntax
     expected by CustomLayoutGenerator — ready to paste into the row-2 cell of the
     audio-script Word table.
  3. Writes a companion "draft" .json file (words, boundariesMs, per-word confidence
     scores) that the WPF review tool loads to pre-fill its waveform markers instead
     of starting from a blank timeline.

USAGE
-----
    python align_batch.py

By default, everything lives in ONE folder called "share", sitting right
next to this script — no arguments needed. Put your audio files and their
matching text/.docx transcripts (same filename stem) together in there:

    align_engine/
      align_batch.py
      share/
        MindsetIELTS_L3_09.mp3
        MindsetIELTS_L3_09.docx      (or .txt — same stem, no markers)
        output/                       <- created automatically
          MindsetIELTS_L3_09_markers.txt
          MindsetIELTS_L3_09_draft.json

If you'd rather keep audio/text/output in separate folders (the old
layout), pass --audio_dir / --text_dir / --out_dir to override any of
them individually — anything you don't override still falls back to
the "share" folder (or "share/output" for the output).

Output:
    <out_dir>/<stem>_markers.txt   <- paste into Word
    <out_dir>/<stem>_draft.json    <- open in the WPF review tool

See README.md in this folder for installation (ffmpeg + pip) and important
notes on the alignment model's license.
"""

import argparse
import json
import re
import sys
from pathlib import Path

import torch

from ctc_forced_aligner import (
    generate_emissions,
    get_alignments,
    get_spans,
    load_alignment_model,
    load_audio,
    postprocess_results,
    preprocess_text,
)

AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}
SCRIPT_EXTS = (".docx", ".txt")  # checked in this order — .docx preferred when both exist


# ---------------------------------------------------------------------------
# Reading the script text: either a plain .txt, or the same two-row Word
# table (AudioTrackName row + script row) that AudioMarkerReview and
# CustomLayoutGenerator both read — so you only maintain ONE source file
# per audio, not a .docx plus a hand-kept-in-sync .txt.
# ---------------------------------------------------------------------------

def read_docx_script_text(docx_path: Path) -> str:
    """Extracts row[1]'s plain text from the first qualifying 2-row table
    (non-empty row[0] = AudioTrackName) in the .docx, exactly mirroring the
    table CustomLayoutGenerator/AudioMarkerReview locate. Only the FIRST
    matching table is used — split multi-track docs into one file per
    track if you need more than that.
    """
    import docx  # python-docx

    document = docx.Document(str(docx_path))
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        track_name = table.rows[0].cells[0].text.strip()
        if not track_name:
            continue
        parts = []
        for cell in table.rows[1].cells:
            for para in cell.paragraphs:
                parts.append(para.text)
        return "\n".join(parts)
    raise ValueError(f"No 2-row table with a non-empty first row found in {docx_path.name}")


def read_script_text(text_path: Path) -> str:
    if text_path.suffix.lower() == ".docx":
        return read_docx_script_text(text_path)
    return text_path.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# Core alignment for a single file pair
# ---------------------------------------------------------------------------

def align_one(model, tokenizer, audio_path: Path, text_path: Path, language: str,
              batch_size: int, window_size: int, context_size: int) -> dict:
    """Run word-level forced alignment for one (audio, text) pair.

    Returns a dict with the ORIGINAL text plus the per-word results in the
    exact order they appear in the text (so index i of `results` always
    corresponds to word i of `original_text.split()`).
    """
    original_text = read_script_text(text_path)
    # The library itself collapses newlines to spaces internally before
    # splitting on whitespace for alignment purposes — we keep our own copy
    # of the untouched original_text so we can restore real line breaks later.
    flat_text = original_text.replace("\n", " ").strip()

    audio_waveform = load_audio(str(audio_path), model.dtype, model.device)
    emissions, stride = generate_emissions(
        model, audio_waveform, window_size, context_size, batch_size
    )

    tokens_starred, text_starred = preprocess_text(
        flat_text, romanize=True, language=language, split_size="word", star_frequency="edges"
    )

    segments, scores, blank_token = get_alignments(emissions, tokens_starred, tokenizer)
    spans = get_spans(tokens_starred, segments, blank_token)
    results = postprocess_results(text_starred, spans, stride, scores)

    # postprocess_results already drops the <star> edge tokens, so `results`
    # should line up 1:1, in order, with original_text's whitespace-split words.
    expected_words = flat_text.split()
    if len(results) != len(expected_words):
        raise ValueError(
            f"Alignment produced {len(results)} word segments but the transcript has "
            f"{len(expected_words)} words for {text_path.name}. This usually means the "
            f"audio and text don't match exactly (missing/extra words, wrong file paired)."
        )

    return {"original_text": original_text, "results": results}


# ---------------------------------------------------------------------------
# Reconstruct the "[ms]word[ms]word..." marker text, preserving the user's
# original whitespace/line breaks exactly, and splitting silence gaps at
# the midpoint between two words (since one marker == shared boundary
# between the end of the previous span and the start of the next).
# ---------------------------------------------------------------------------

def build_marker_text(original_text: str, results: list) -> tuple[str, list, list, list]:
    tokens = re.findall(r"\S+|\s+", original_text)
    n = len(results)
    if n == 0:
        raise ValueError("No aligned words to build markers from.")

    boundaries_ms = [0] * (n + 1)
    boundaries_ms[0] = round(results[0]["start"] * 1000)
    for i in range(1, n):
        boundaries_ms[i] = round(((results[i - 1]["end"] + results[i]["start"]) / 2) * 1000)
    boundaries_ms[n] = round(results[-1]["end"] * 1000)

    words = [r["text"] for r in results]
    scores = [round(r["score"], 4) for r in results]

    out = []
    word_idx = 0
    for tok in tokens:
        if tok.strip() == "":
            out.append(tok)
            continue
        if word_idx == 0:
            out.append(f"[{boundaries_ms[0]}]")
        out.append(tok)
        out.append(f"[{boundaries_ms[word_idx + 1]}]")
        word_idx += 1

    return "".join(out), words, boundaries_ms, scores


def collapse_to_sentence_boundaries(words, boundaries_ms):
    """Optional: keep only markers at sentence-ending punctuation, for a
    coarser (sentence-level) marker file instead of word-level."""
    sentence_end = re.compile(r"[.!?]\"?'?$")
    keep = {0, len(words)}
    for i, w in enumerate(words):
        if sentence_end.search(w):
            keep.add(i + 1)
    return sorted(keep)


# ---------------------------------------------------------------------------
# Batch driver
# ---------------------------------------------------------------------------

def find_pairs(audio_dir: Path, text_dir: Path):
    pairs = []
    for audio_path in sorted(audio_dir.iterdir()):
        if audio_path.suffix.lower() not in AUDIO_EXTS:
            continue
        text_path = None
        for ext in SCRIPT_EXTS:
            candidate = text_dir / f"{audio_path.stem}{ext}"
            if candidate.exists():
                text_path = candidate
                break
        if text_path is not None:
            pairs.append((audio_path, text_path))
        else:
            print(f"  [skip] no matching .docx or .txt for {audio_path.name} "
                  f"(expected {audio_path.stem}.docx or {audio_path.stem}.txt)", file=sys.stderr)
    return pairs


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--share_dir", type=Path, default=None,
                     help="Single folder holding BOTH audio and text/.docx files, matched by "
                          "filename stem. Defaults to a 'share' folder next to this script. "
                          "Outputs go to <share_dir>/output unless --out_dir overrides it.")
    ap.add_argument("--audio_dir", type=Path, default=None,
                     help="Override: use a separate folder for audio instead of --share_dir.")
    ap.add_argument("--text_dir", type=Path, default=None,
                     help="Override: use a separate folder for text/.docx instead of --share_dir.")
    ap.add_argument("--out_dir", type=Path, default=None,
                     help="Override: write outputs somewhere other than <share_dir>/output.")
    ap.add_argument("--language", default="eng", help="ISO 639-3 code (default: eng)")
    ap.add_argument("--granularity", choices=["word", "sentence"], default="word",
                     help="word = a marker after every word (default); "
                          "sentence = markers only at sentence-ending punctuation")
    ap.add_argument("--device", default="cuda" if torch.cuda.is_available() else "cpu")
    ap.add_argument("--batch_size", type=int, default=8)
    ap.add_argument("--window_size", type=int, default=30)
    ap.add_argument("--context_size", type=int, default=2)
    ap.add_argument("--alignment_model", default="MahmoudAshraf/mms-300m-1130-forced-aligner",
                     help="HF model id or local path. NOTE: the default model is "
                          "CC-BY-NC-4.0 licensed (non-commercial) — see README.md")
    args = ap.parse_args()

    script_dir = Path(__file__).resolve().parent
    share_dir = args.share_dir or (script_dir / "share")
    audio_dir = args.audio_dir or share_dir
    text_dir = args.text_dir or share_dir
    out_dir = args.out_dir or (share_dir / "output")

    for label, d in (("audio", audio_dir), ("text", text_dir)):
        if not d.exists():
            print(f"{label.capitalize()} folder not found: {d}\n"
                  f"Create it and put your audio + matching text/.docx files inside "
                  f"(same filename stem for each pair), then re-run.", file=sys.stderr)
            sys.exit(1)

    out_dir.mkdir(parents=True, exist_ok=True)

    pairs = find_pairs(audio_dir, text_dir)
    if not pairs:
        print("No matching (audio, text) pairs found.", file=sys.stderr)
        sys.exit(1)
    print(f"Found {len(pairs)} file pair(s). Loading alignment model on {args.device} ...")

    dtype = torch.float16 if args.device == "cuda" else torch.float32
    model, tokenizer = load_alignment_model(args.device, args.alignment_model, dtype=dtype)

    ok, failed = 0, []
    for audio_path, text_path in pairs:
        print(f"Aligning {audio_path.name} ...")
        try:
            aligned = align_one(
                model, tokenizer, audio_path, text_path, args.language,
                args.batch_size, args.window_size, args.context_size,
            )
            marker_text, words, boundaries_ms, scores = build_marker_text(
                aligned["original_text"], aligned["results"]
            )

            if args.granularity == "sentence":
                keep_idx = collapse_to_sentence_boundaries(words, boundaries_ms)
                # Re-flatten with only the kept boundary markers. Each boundary is
                # emitted exactly once, right after the word that precedes it — it
                # simultaneously serves as the start marker of the next chunk, so
                # we must NOT also emit it again before that next word.
                tokens = re.findall(r"\S+|\s+", aligned["original_text"])
                out, wi = [], 0
                for tok in tokens:
                    if tok.strip() == "":
                        out.append(tok)
                        continue
                    if wi == 0:
                        out.append(f"[{boundaries_ms[0]}]")
                    out.append(tok)
                    if (wi + 1) in keep_idx:
                        out.append(f"[{boundaries_ms[wi + 1]}]")
                    wi += 1
                marker_text = "".join(out)

            stem = audio_path.stem
            (out_dir / f"{stem}_markers.txt").write_text(marker_text, encoding="utf-8")
            draft = {
                "audioFile": audio_path.name,
                "words": words,
                "boundariesMs": boundaries_ms,
                "scores": scores,
            }
            (out_dir / f"{stem}_draft.json").write_text(
                json.dumps(draft, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            low_conf = sum(1 for s in scores if s < -1.5)
            print(f"  -> wrote {stem}_markers.txt and {stem}_draft.json "
                  f"({len(words)} words, {low_conf} low-confidence)")
            ok += 1
        except Exception as e:  # noqa: BLE001 — batch job, keep going on per-file failure
            print(f"  [FAILED] {audio_path.name}: {e}", file=sys.stderr)
            failed.append(audio_path.name)

    print(f"\nDone: {ok} succeeded, {len(failed)} failed.")
    if failed:
        print("Failed files:", ", ".join(failed))
        sys.exit(1)


if __name__ == "__main__":
    main()
